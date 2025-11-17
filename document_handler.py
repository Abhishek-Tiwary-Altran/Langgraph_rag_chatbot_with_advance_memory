import os
from typing import List
import PyPDF2
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangchainDocument

# os.environ["OTEL_EXPORTER_OTLP_ENDPOINT"] = f"https://aps-workspaces.us-east-1.amazonaws.com/v1/metrics"

class DocumentHandler:
    def __init__(self, persist_directory: str = "./chroma_db"):
        self.persist_directory = persist_directory
        self.text_splitter = RecursiveCharacterTextSplitter.from_tiktoken_encoder(
            chunk_size=512, 
            chunk_overlap=0
        )

    def read_pdf(self, file) -> str:
        """Read PDF file and return text content."""
        pdf_reader = PyPDF2.PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
        return text

    def read_docx(self, file) -> str:
        """Read DOCX file and return text content."""
        doc = Document(file)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text

    def read_txt(self, file) -> str:
        """Read TXT file and return text content."""
        return file.read().decode('utf-8')

    def process_file(self, file) -> List[LangchainDocument]:
        """Process uploaded file and return list of documents."""
        file_name = file.name
        file_extension = os.path.splitext(file_name)[1].lower()

        if file_extension == '.pdf':
            text = self.read_pdf(file)
        elif file_extension == '.docx':
            text = self.read_docx(file)
        elif file_extension == '.txt':
            text = self.read_txt(file)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")

        # Create a Langchain Document
        doc = LangchainDocument(
            page_content=text,
            metadata={"source": file_name}
        )

        # Split the document
        return self.text_splitter.split_documents([doc])
