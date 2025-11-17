import os
from typing import List, Optional
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import WebBaseLoader
from langchain_community.vectorstores import Chroma
from langchain.schema import Document as LangchainDocument
import PyPDF2
from docx import Document
import logging

# os.environ["OTEL_EXPORTER_OTLP_ENDPOINT"] = f"https://aps-workspaces.us-east-1.amazonaws.com/v1/metrics"

# Configure logging
logger = logging.getLogger(__name__)

def process_file(file) -> List[LangchainDocument]:
    """Process uploaded file and return list of documents.
    
    Args:
        file: Streamlit uploaded file object
        
    Returns:
        List of document chunks
    """
    file_name = file.name
    file_extension = os.path.splitext(file_name)[1].lower()
    
    logger.info(f"Processing file: {file_name} with extension {file_extension}")
    
    # Read file content based on file type
    if file_extension == '.pdf':
        text = read_pdf(file)
    elif file_extension == '.docx':
        text = read_docx(file)
    elif file_extension == '.txt':
        text = read_txt(file)
    else:
        raise ValueError(f"Unsupported file type: {file_extension}")
    
    # Create a Langchain Document
    doc = LangchainDocument(
        page_content=text,
        metadata={"source": file_name}
    )
    
    # Split the document
    text_splitter = RecursiveCharacterTextSplitter.from_tiktoken_encoder(
        chunk_size=512, 
        chunk_overlap=50
    )
    doc_splits = text_splitter.split_documents([doc])
    logger.info(f"Split document into {len(doc_splits)} chunks")
    
    return doc_splits

def read_pdf(file) -> str:
    """Read PDF file and return text content."""
    pdf_reader = PyPDF2.PdfReader(file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text() + "\n"
    return text

def read_docx(file) -> str:
    """Read DOCX file and return text content."""
    doc = Document(file)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

def read_txt(file) -> str:
    """Read TXT file and return text content."""
    return file.read().decode('utf-8')

def create_vectorstore(doc_splits, embed_model, collection_name="user-documents"):
    """Create a vector store from document splits."""
    persist_dir = "./chroma_db"
    return Chroma.from_documents(
        documents=doc_splits, 
        embedding=embed_model, 
        collection_name=collection_name,
        persist_directory=persist_dir
    )

def create_retriever(vectorstore):
    """Create a retriever from vector store."""
    return vectorstore.as_retriever(search_kwargs={"k": 5})

def add_documents_to_vectorstore(vectorstore, doc_splits):
    """Add documents to an existing vector store.
    
    Args:
        vectorstore: Existing Chroma vector store
        doc_splits: Document chunks to add
        
    Returns:
        Updated vector store
    """
    if not doc_splits:
        return vectorstore
        
    logger.info(f"Adding {len(doc_splits)} document chunks to vector store")
    vectorstore.add_documents(doc_splits)
    vectorstore.persist()
    logger.info("Documents added and vector store persisted")
    return vectorstore